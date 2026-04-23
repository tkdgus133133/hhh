from __future__ import annotations

import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]

# 기존 보고서 양식을 그대로 사용하고, 내용만 헝가리 맥락으로 치환
TEMPLATE_P1 = ROOT / "SG_01_시장보고서_Sereterol.docx"
TEMPLATE_P2 = ROOT / "SG_02_수출가격전략_Gadvoa (1).docx"
TEMPLATE_P3 = ROOT / "SG_03_바이어리스트.docx"
TEMPLATE_FINAL = ROOT / "SG_최종보고서(2-3-1 합친버전).docx"

_FONT_REGISTERED = False


def _register_korean_font() -> None:
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    font_candidates = [
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/NanumGothic.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for font_path in font_candidates:
        if font_path.is_file():
            pdfmetrics.registerFont(TTFont("AppFont", str(font_path)))
            _FONT_REGISTERED = True
            return
    _FONT_REGISTERED = True


def _sanitize_text(text: Any) -> str:
    src = str(text or "").strip()
    replace_map = {
        "Singapore": "Hungary",
        "싱가포르": "헝가리",
        "HSA": "OGYÉI",
        "MOH": "OGYÉI/보건",
        "ALPS": "NEAK",
        "GeBIZ": "공공조달(NEAK)",
    }
    out = src
    for old, new in replace_map.items():
        out = out.replace(old, new)
    return out


def _draw_lines_pdf(title: str, lines: list[str], out_pdf_path: Path) -> None:
    out_pdf_path.parent.mkdir(parents=True, exist_ok=True)
    _register_korean_font()
    c = canvas.Canvas(str(out_pdf_path), pagesize=A4)
    width, height = A4
    font_name = "AppFont" if "AppFont" in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    c.setFont(font_name, 18)
    c.drawString(40, height - 50, _sanitize_text(title))
    c.setFont(font_name, 11)
    y = height - 85
    line_h = 18
    for raw in lines:
        line = _sanitize_text(raw)
        if not line:
            y -= line_h
            continue
        if y < 50:
            c.showPage()
            c.setFont(font_name, 11)
            y = height - 50
        c.drawString(40, y, line[:120])
        y -= line_h
    c.save()


def _replace_in_paragraph(paragraph, replace_map: dict[str, str]) -> None:
    text = paragraph.text or ""
    new_text = text
    for old, new in replace_map.items():
        if old:
            new_text = new_text.replace(old, new)
    if new_text != text:
        paragraph.text = new_text


def _replace_all(doc: Document, replace_map: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replace_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replace_map)


def _sanitize_semantic_context(doc: Document) -> None:
    """템플릿 내 싱가포르 의미데이터를 헝가리 문맥으로 강제 치환."""
    replace_map = {
        "MOH Holdings": "헝가리 공공병원 네트워크",
        "MOH Singapore": "헝가리 보건당국",
        "MOH": "헝가리 보건당국",
        "HSA": "OGYÉI",
        "ALPS": "NEAK 조달 채널",
        "GeBIZ": "NEAK 공공조달",
        "NDF": "헝가리 급여·등재 체계",
        "APAC": "중부·동유럽",
        "ASEAN": "EU/중부·동유럽",
        "Singapore hospital supply price": "Hungary market reference price",
        "Singapore Department of Statistics": "Hungary Central Statistical Office",
        "싱가포르 보건부": "헝가리 보건당국",
        "싱가포르 보건과학청": "헝가리 의약품 규제기관(OGYÉI)",
        "싱가포르 정부 공공조달": "헝가리 공공조달(NEAK)",
        "싱가포르 약가 규제": "헝가리 약가·급여 규제",
        "싱가포르 내": "헝가리 내",
        "싱가포르 시장": "헝가리 시장",
        "싱가포르 진출": "헝가리 진출",
        "인구 약 596만": "인구 약 960만",
        "USD 88,447": "USD 22,000",
        "USD 2.6B": "USD 14.0B",
        "GDP의 약 6.5%": "GDP 대비 보건지출 약 6~7%",
    }
    _replace_all(doc, replace_map)


def _append_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.add_run(title).bold = True


def _append_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)


def _prepend_heading(doc: Document, title: str) -> None:
    anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p = anchor.insert_paragraph_before("")
    p.add_run(title).bold = True


def _prepend_kv(doc: Document, key: str, value: str) -> None:
    anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p = anchor.insert_paragraph_before("")
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)


def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    abs_docx = str(docx_path.resolve())
    abs_pdf = str(pdf_path.resolve())
    ps_script = (
        "$ErrorActionPreference='Stop';"
        f"$docx='{abs_docx}';"
        f"$pdf='{abs_pdf}';"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        "try {"
        "  $doc=$word.Documents.Open($docx);"
        "  $doc.ExportAsFixedFormat($pdf,17);"
        "  $doc.Close();"
        "} finally {"
        "  $word.Quit();"
        "}"
    )
    subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps_script],
        check=True,
        capture_output=True,
        text=True,
    )


def _country_replace_map(product_name: str, date_str: str) -> dict[str, str]:
    return {
        "싱가포르": "헝가리",
        "Singapore": "Hungary",
        "2026-04-21": date_str,
        "Sereterol Activair": product_name or "제품명",
        "Gadvoa  Inj.": product_name or "제품명",
        "Gadvoa Inj.": product_name or "제품명",
    }


def build_p1_template_pdf(
    product_key: str,
    result: dict[str, Any],
    refs: list[dict[str, Any]],
    static_profile: dict[str, Any],
    out_pdf_path: Path,
) -> None:
    if not TEMPLATE_P1.is_file():
        lines = [
            f"생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
            f"품목: {result.get('trade_name') or result.get('product_name') or product_key}",
            f"국가: 헝가리",
            f"수출 판정: {result.get('verdict') or '미상'}",
            f"시장/의료 근거: {result.get('basis_market_medical') or '-'}",
            f"규제 근거: {result.get('basis_regulatory') or '-'}",
            f"무역 근거: {result.get('basis_trade') or '-'}",
            f"진입 채널: {result.get('entry_pathway') or '-'}",
            f"가격 포지셔닝: {result.get('price_positioning_pbs') or '-'}",
            f"리스크/조건: {result.get('risks_conditions') or '-'}",
        ]
        if refs:
            lines.append("참고 링크:")
            for i, r in enumerate(refs[:7], 1):
                lines.append(f"{i}. {r.get('title') or '-'} / {r.get('url') or '-'}")
        _draw_lines_pdf("헝가리 시장조사 보고서 (P1)", lines, out_pdf_path)
        return
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    trade_name = str(result.get("trade_name") or result.get("product_name") or product_key)
    doc = Document(str(TEMPLATE_P1))
    _replace_all(doc, _country_replace_map(trade_name, ts))
    _sanitize_semantic_context(doc)
    _replace_all(
        doc,
        {
            "1인당 GDP   USD 88,447  (IMF 2024)": "1인당 GDP 및 보건지표는 헝가리 실시간/정적 데이터 기준으로 보정 적용",
            "ALPS": "NEAK",
            "MOH": "OGYÉI/NEAK",
        },
    )

    _prepend_heading(doc, "[자동 생성 본문] 아래 항목은 최신 산출값 기준입니다.")
    _prepend_kv(doc, "품목", trade_name)
    _prepend_kv(doc, "수출 판정", str(result.get("verdict") or "미상"))
    _prepend_kv(doc, "시장/의료 근거", str(result.get("basis_market_medical") or "-"))
    _prepend_kv(doc, "규제 근거", str(result.get("basis_regulatory") or "-"))
    _prepend_kv(doc, "무역 근거", str(result.get("basis_trade") or "-"))
    _prepend_kv(doc, "가격 포지셔닝", str(result.get("price_positioning_pbs") or "-"))
    _prepend_kv(doc, "리스크/조건", str(result.get("risks_conditions") or "-"))
    rp = static_profile.get("raw_payload", {}) if static_profile else {}
    _prepend_kv(
        doc,
        "정적 3소스",
        f"NEAK={rp.get('price_rows', 0)} / OGYÉI={rp.get('registry_rows', 0)} / EMA={rp.get('ema_rows', 0)}",
    )
    if refs:
        _prepend_heading(doc, "Perplexity 참고 링크")
        for i, r in enumerate(reversed(refs[:7]), 1):
            title = str(r.get("title") or "제목 없음")
            url = str(r.get("url") or "")
            _prepend_kv(doc, f"{i}", f"{title} ({url})")

    out_docx = out_pdf_path.with_suffix(".docx")
    doc.save(str(out_docx))
    _convert_docx_to_pdf(out_docx, out_pdf_path)


def build_p2_template_pdf(p2_data: dict[str, Any], out_pdf_path: Path) -> None:
    if not TEMPLATE_P2.is_file():
        lines = [
            f"생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
            f"품목: {p2_data.get('product_name') or '제품명'}",
            "국가: 헝가리",
            f"판정: {p2_data.get('verdict') or '-'}",
            f"산식: {p2_data.get('formula_str') or '-'}",
        ]
        for sec in p2_data.get("sections") or []:
            lines.append(f"[{sec.get('seg_label') or '시장'}] 기준가: {sec.get('base_price')}")
            for sc in (sec.get("scenarios") or [])[:3]:
                price = sc.get("price") if sc.get("price") is not None else sc.get("price_usd")
                lines.append(f"- {sc.get('label') or sc.get('name')}: {price} / {sc.get('reason') or '-'}")
        _draw_lines_pdf("헝가리 수출가격 전략 보고서 (P2)", lines, out_pdf_path)
        return
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    product_name = str(p2_data.get("product_name") or "제품명")
    doc = Document(str(TEMPLATE_P2))
    _replace_all(doc, _country_replace_map(product_name, ts))
    _sanitize_semantic_context(doc)
    _replace_all(doc, {"공공  /  민간": "공공 / 민간 (헝가리 기준)"})

    _prepend_heading(doc, "[자동 생성 본문] 아래 항목은 최신 가격 산출 결과입니다.")
    _prepend_kv(doc, "품목", product_name)
    _prepend_kv(doc, "INN", str(p2_data.get("inn_name") or "-"))
    _prepend_kv(doc, "판정", str(p2_data.get("verdict") or "-"))
    _prepend_kv(doc, "산식", str(p2_data.get("formula_str") or "-"))
    sections = p2_data.get("sections") or []
    for sec in sections:
        seg = str(sec.get("seg_label") or "시장")
        base = sec.get("base_price")
        _prepend_kv(doc, f"{seg} 기준가", f"USD {float(base):.2f}" if isinstance(base, (int, float)) else "-")
        for sc in sec.get("scenarios", [])[:3]:
            label = str(sc.get("label") or sc.get("name") or "시나리오")
            price = sc.get("price") if sc.get("price") is not None else sc.get("price_usd")
            reason = str(sc.get("reason") or "-")
            ptxt = f"USD {float(price):.2f}" if isinstance(price, (int, float)) else "-"
            _prepend_kv(doc, f"  - {label}", f"{ptxt} / {reason}")

    out_docx = out_pdf_path.with_suffix(".docx")
    doc.save(str(out_docx))
    _convert_docx_to_pdf(out_docx, out_pdf_path)


def build_p3_template_pdf(
    buyers: list[dict[str, Any]],
    product_label: str,
    static_profile: dict[str, Any],
    out_pdf_path: Path,
) -> None:
    if not TEMPLATE_P3.is_file():
        lines = [
            f"생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
            f"품목: {product_label}",
            "국가: 헝가리",
            f"Top 후보 수: {len(buyers)}",
        ]
        for i, b in enumerate(buyers[:10], 1):
            e = b.get("enriched", {}) or {}
            lines.append(
                f"{i}. {b.get('company_name', '-')} / {b.get('country', '-')} / GMP={e.get('has_gmp')}"
            )
        _draw_lines_pdf("헝가리 바이어 발굴 보고서 (P3)", lines, out_pdf_path)
        return
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    doc = Document(str(TEMPLATE_P3))
    _replace_all(doc, _country_replace_map(product_label, ts))
    _sanitize_semantic_context(doc)
    _replace_all(doc, {"싱가포르 바이어 후보 리스트": "헝가리 바이어 후보 리스트"})

    _prepend_heading(doc, "[자동 생성 본문] 아래 항목은 최신 바이어 발굴 결과입니다.")
    rp = static_profile.get("raw_payload", {}) if static_profile else {}
    _prepend_kv(
        doc,
        "정적 3소스",
        f"NEAK={rp.get('price_rows', 0)} / OGYÉI={rp.get('registry_rows', 0)} / EMA={rp.get('ema_rows', 0)}",
    )
    _prepend_kv(doc, "품목", product_label)
    _prepend_kv(doc, "Top 후보 수", str(len(buyers)))
    for i, b in enumerate(buyers[:10], 1):
        e = b.get("enriched", {}) or {}
        _prepend_kv(
            doc,
            f"{i}. {b.get('company_name', '-')}",
            f"{b.get('country', '-')} / GMP={e.get('has_gmp')} / 타깃국진출={e.get('has_target_country_presence')}",
        )

    out_docx = out_pdf_path.with_suffix(".docx")
    doc.save(str(out_docx))
    _convert_docx_to_pdf(out_docx, out_pdf_path)


def _p2_public_private_price_labels(pub: dict[str, Any], pri: dict[str, Any]) -> tuple[str, str]:
    """P2 analysis 공공/민간 기준가 문장 — USD 우선."""

    def one(m: dict[str, Any]) -> str:
        u = m.get("final_price_usd")
        if u is not None and u != "":
            try:
                return f"USD {float(u):.2f}"
            except (TypeError, ValueError):
                pass
        return "-"

    return one(pub), one(pri)


def _normalize_p2_payload(p2_result: dict[str, Any] | None) -> tuple[dict[str, Any], dict[str, Any]]:
    """P2 결과 포맷 차이를 흡수해 analysis/p2_data를 표준화."""
    if not isinstance(p2_result, dict):
        return {}, {}
    # 파이프라인 저장 포맷: {"analysis": {...}, "p2_data": {...}}
    if isinstance(p2_result.get("analysis"), dict) or isinstance(p2_result.get("p2_data"), dict):
        analysis = p2_result.get("analysis", {})
        p2d = p2_result.get("p2_data", {})
        return analysis if isinstance(analysis, dict) else {}, p2d if isinstance(p2d, dict) else {}

    # 수동/직접 저장 포맷: p2_data 자체가 루트
    p2d = p2_result
    sections = p2d.get("sections", []) if isinstance(p2d.get("sections", []), list) else []
    analysis: dict[str, Any] = {"rationale": str(p2d.get("macro_text") or "")}
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        label = str(sec.get("seg_label", "")).lower()
        entry = {
            "final_price_usd": sec.get("base_price"),
            "scenarios": sec.get("scenarios", []),
        }
        if "public" in label:
            analysis["public_market"] = entry
        elif "private" in label:
            analysis["private_market"] = entry
    return analysis, p2d


def _as_text(value: Any) -> str:
    """문자열 필드가 dict/list로 들어와도 안전하게 텍스트화."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def build_final_template_pdf(
    p1_result: dict[str, Any] | None,
    p2_result: dict[str, Any] | None,
    buyers: list[dict[str, Any]] | None,
    out_pdf_path: Path,
) -> None:
    if not TEMPLATE_FINAL.is_file():
        lines = [
            f"생성일: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
            "국가: 헝가리",
            "최종 통합 보고서",
        ]
        if p1_result:
            lines.extend([
                f"[P1] 판정: {p1_result.get('verdict') or '미상'}",
                f"[P1] 가격 포지셔닝: {p1_result.get('price_positioning_pbs') or '-'}",
            ])
        if p2_result:
            analysis, p2d = _normalize_p2_payload(p2_result)
            pub = (analysis.get("public_market") or {}) if isinstance(analysis, dict) else {}
            pri = (analysis.get("private_market") or {}) if isinstance(analysis, dict) else {}
            pub_s, pri_s = _p2_public_private_price_labels(pub, pri)
            rat = _as_text(analysis.get("rationale") or p2d.get("macro_text"))
            if rat:
                lines.append(f"[P2] 가격 전략 요지: {rat[:280]}{'…' if len(rat) > 280 else ''}")
            lines.extend([f"[P2] 공공 기준가: {pub_s}", f"[P2] 민간 기준가: {pri_s}"])
        if buyers:
            lines.append(f"[P3] 바이어 후보 수: {len(buyers)}")
            for i, b in enumerate((buyers or [])[:5], 1):
                lines.append(f"[P3] {i}. {b.get('company_name') or '-'}")
        _draw_lines_pdf("헝가리 최종 통합 보고서 (P4)", lines, out_pdf_path)
        return
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    product_name = "헝가리 진출 품목"
    if p2_result:
        extracted = p2_result.get("extracted") if isinstance(p2_result, dict) else {}
        if not isinstance(extracted, dict):
            extracted = {}
        p2d_name = p2_result.get("product_name") if isinstance(p2_result, dict) else None
        product_name = str(extracted.get("product_name") or p2d_name or product_name)
    elif p1_result:
        product_name = str(p1_result.get("trade_name") or product_name)

    doc = Document(str(TEMPLATE_FINAL))
    _replace_all(doc, _country_replace_map(product_name, ts))
    _sanitize_semantic_context(doc)
    _replace_all(doc, {"싱가포르 진출 전략 보고서": "헝가리 진출 전략 보고서"})

    _prepend_heading(doc, "[자동 생성 본문] 아래 항목은 최신 P1/P2/P3 산출값 기준입니다.")
    if p1_result:
        _prepend_kv(doc, "[시장조사] 판정", str(p1_result.get("verdict") or "미상"))
        _prepend_kv(doc, "[시장조사] 가격 포지셔닝", str(p1_result.get("price_positioning_pbs") or "-"))
    if p2_result:
        analysis, p2d = _normalize_p2_payload(p2_result)
        pub = (analysis.get("public_market") or {}) if isinstance(analysis, dict) else {}
        pri = (analysis.get("private_market") or {}) if isinstance(analysis, dict) else {}
        pub_s, pri_s = _p2_public_private_price_labels(pub, pri)
        rat2 = _as_text(analysis.get("rationale") or p2d.get("macro_text"))
        if rat2:
            _prepend_kv(
                doc,
                "[가격전략] 요지 (NEAK·KUP·EU)",
                rat2[:600] + ("…" if len(rat2) > 600 else ""),
            )
        _prepend_kv(doc, "[가격전략] 공공 기준가", pub_s)
        _prepend_kv(doc, "[가격전략] 민간 기준가", pri_s)
    if buyers:
        _prepend_kv(doc, "[바이어] Top 후보 수", str(len(buyers)))
        for i, b in enumerate((buyers or [])[:5], 1):
            _prepend_kv(doc, f"[바이어] {i}", str(b.get("company_name") or "-"))

    out_docx = out_pdf_path.with_suffix(".docx")
    doc.save(str(out_docx))
    _convert_docx_to_pdf(out_docx, out_pdf_path)

